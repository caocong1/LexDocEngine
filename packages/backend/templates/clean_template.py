#!/usr/bin/env python3
"""
法律备忘录模板清理脚本
自动清理模板中的示例内容，并添加 Carbone 占位符
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def clean_template(input_path, output_path):
    """
    清理模板并添加 Carbone 占位符

    Args:
        input_path: 原始模板路径
        output_path: 输出模板路径
    """
    print(f"正在读取模板: {input_path}")
    doc = Document(input_path)

    # 需要删除的文本模式
    texts_to_remove = [
        "中翼航空投资有限公司",
        "旅客投诉餐食异物相关事宜",
        "国浩律师（北京）事务所",
        "林大荣",
        "香港籍",
        "顺义区医院",
        "2026年1月8日",
        "2026年1月19日",
    ]

    # 遍历所有段落
    paragraphs_to_process = []

    for para in doc.paragraphs:
        text = para.text.strip()

        # 跳过空段落
        if not text:
            continue

        print(f"处理段落: {text[:50]}...")

        # 检查是否包含需要删除的文本
        contains_example = any(pattern in text for pattern in texts_to_remove)

        if contains_example:
            # 标记需要处理的段落
            paragraphs_to_process.append((para, text))

    # 创建替换规则
    replacements = {
        # 标题部分
        r"致：.*公司": "致：{d.client_name}",
        r"关于.*之法律备忘录": "关于 {d.case_title} 之法律备忘录",

        # 开场白
        r"我们，.*律师.*事务所.*，接受.*委托": "我们接受委托",
        r"就.*相关事宜": "就相关事宜",

        # 日期
        r"\d{4}年\d{1,2}月\d{1,2}日": "{d.event_date}",
        r"二〇二\d年.*月.*日": "{d.date}",

        # 具体人名/公司名
        r"贵司": "委托方",
    }

    # 应用替换
    for para, original_text in paragraphs_to_process:
        new_text = original_text
        for pattern, replacement in replacements.items():
            new_text = re.sub(pattern, replacement, new_text)

        # 如果文本改变了，更新段落
        if new_text != original_text:
            para.text = new_text
            print(f"  替换: {original_text[:40]} -> {new_text[:40]}")

    # 查找"一、基本事实"章节并替换内容
    found_basic_facts = False
    found_legal_analysis = False
    found_recommendations = False

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # 找到"一、基本事实"
        if "一、基本事实" in text or "基本事实" in text:
            found_basic_facts = True
            print(f"\n找到'基本事实'章节，索引: {i}")

            # 删除后续段落直到下一章节
            j = i + 1
            while j < len(doc.paragraphs):
                next_para = doc.paragraphs[j]
                next_text = next_para.text.strip()

                # 如果遇到下一章节，停止
                if re.match(r"[（(]?[一二三四五六七八九十]+[)）]?、", next_text) or \
                   re.match(r"\d+[、.]", next_text):
                    break

                # 删除这个段落的内容（保留段落用于添加占位符）
                if j == i + 1:
                    # 第一个段落保留用于添加占位符
                    next_para.text = "{d.basic_facts}"
                    print(f"  添加占位符: {{d.basic_facts}}")
                    j += 1
                    break
                else:
                    # 其他段落清空
                    next_para.text = ""

                j += 1

        # 找到"二、法律分析"
        elif "法律分析" in text:
            found_legal_analysis = True
            print(f"\n找到'法律分析'章节，索引: {i}")

            j = i + 1
            while j < len(doc.paragraphs):
                next_para = doc.paragraphs[j]
                next_text = next_para.text.strip()

                if re.match(r"[（(]?[一二三四五六七八九十]+[)）]?、", next_text) or \
                   re.match(r"\d+[、.]", next_text):
                    break

                if j == i + 1:
                    next_para.text = "{d.legal_opinion}"
                    print(f"  添加占位符: {{d.legal_opinion}}")
                    j += 1
                    break
                else:
                    next_para.text = ""

                j += 1

        # 找到"三、后续建议"或"建议"
        elif "建议" in text and "后续" in text:
            found_recommendations = True
            print(f"\n找到'后续建议'章节，索引: {i}")

            j = i + 1
            while j < len(doc.paragraphs):
                next_para = doc.paragraphs[j]
                next_text = next_para.text.strip()

                # 如果是落款（律所名称和日期），停止
                if "律师" in next_text or "事务所" in next_text:
                    # 替换落款
                    next_para.text = "{d.law_firm}"
                    print(f"  替换律所名称: {{d.law_firm}}")

                    # 下一行应该是日期
                    if j + 1 < len(doc.paragraphs):
                        doc.paragraphs[j + 1].text = "{d.date}"
                        print(f"  替换日期: {{d.date}}")
                    break

                if j == i + 1:
                    next_para.text = "{d.recommendations}"
                    print(f"  添加占位符: {{d.recommendations}}")
                    j += 1
                    break
                else:
                    next_para.text = ""

                j += 1

    # 保存处理后的模板
    print(f"\n保存清理后的模板到: {output_path}")
    doc.save(output_path)

    print("\n✅ 模板清理完成！")
    print(f"\n已添加的 Carbone 占位符:")
    print("  - {{d.client_name}}        客户名称")
    print("  - {{d.case_title}}         案件标题")
    print("  - {{d.basic_facts}}        基本事实")
    print("  - {{d.legal_opinion}}      法律分析")
    print("  - {{d.recommendations}}    后续建议")
    print("  - {{d.law_firm}}           律所名称")
    print("  - {{d.date}}               日期")

if __name__ == "__main__":
    import sys
    import os

    # 模板路径
    template_dir = os.path.dirname(os.path.abspath(__file__))
    input_template = os.path.join(template_dir, "legal-memo-template.dotx")
    output_template = os.path.join(template_dir, "legal-memo-carbone.dotx")

    if not os.path.exists(input_template):
        print(f"❌ 错误: 找不到模板文件 {input_template}")
        sys.exit(1)

    print("=" * 60)
    print("法律备忘录模板自动清理工具")
    print("=" * 60)

    clean_template(input_template, output_template)

    print("\n" + "=" * 60)
    print(f"✅ 完成！新模板已保存为: {output_template}")
    print("=" * 60)
